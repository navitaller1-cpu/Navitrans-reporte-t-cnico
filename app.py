from flask import Flask, render_template, request, send_from_directory, redirect, url_for, flash
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import os, datetime, json

# Config
app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "change-this-in-production-12345")
app.config["UPLOAD_FOLDER"] = os.path.join(app.root_path, "uploads")
app.config["GENERATED_FOLDER"] = os.path.join(app.root_path, "generated")
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024  # 32 MB

ALLOWED_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp"}

def ensure_dirs():
    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
    os.makedirs(app.config["GENERATED_FOLDER"], exist_ok=True)
    os.makedirs(os.path.join(app.root_path, "static", "img"), exist_ok=True)

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS

def set_font_cambria_10(run):
    """Helper para aplicar fuente Cambria 10pt a un run"""
    run.font.name = 'Cambria'
    run.font.size = Pt(10)

def set_font_cambria_size(run, size_pt):
    """Helper para aplicar fuente Cambria con tamaño específico"""
    run.font.name = 'Cambria'
    run.font.size = Pt(size_pt)

def add_cell_borders(cell):
    """Helper para agregar bordes a una celda"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_thick_outer_borders_only(table):
    """Helper para agregar solo bordes exteriores gruesos a una tabla"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Crear elemento de bordes de tabla
    tblBorders = OxmlElement('w:tblBorders')
    
    # Solo bordes exteriores gruesos
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # Borde grueso
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    # Sin bordes internos
    for border_name in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def set_cell_vertical_alignment(cell, alignment='center'):
    """Helper para alinear verticalmente el contenido de una celda"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    tcPr.append(vAlign)

def merge_cells_horizontal(table, row_idx, start_col, end_col):
    """Helper para combinar celdas horizontalmente"""
    row = table.rows[row_idx]
    cell = row.cells[start_col]
    
    for col_idx in range(start_col + 1, end_col + 1):
        if col_idx < len(row.cells):
            cell.merge(row.cells[col_idx])
    
    return cell

@app.route("/")
def index():
    hoy = datetime.date.today().isoformat()
    return render_template("form.html", today=hoy)

@app.route("/submit", methods=["POST"])
def submit():
    ensure_dirs()
    fecha = request.form.get("fecha", "").strip()
    cliente = request.form.get("cliente", "").strip()
    equipo = request.form.get("equipo", "").strip()
    kilometraje = request.form.get("kilometraje", "").strip()
    horas = request.form.get("horas", "").strip()

    conditions_json = request.form.get("conditions_json", "[]")
    try:
        condiciones = [c for c in (json.loads(conditions_json) or []) if c.get("text")]
    except Exception:
        condiciones = []

    correcciones_descripciones = request.form.getlist("corrections_desc[]")
    correcciones_titulos = request.form.getlist("corrections_title[]")
    correcciones_imagenes = request.files.getlist("corrections_img[]")

    saved_correcciones = []
    for idx, desc in enumerate(correcciones_descripciones):
        desc = (desc or "").strip()
        titulo = correcciones_titulos[idx] if idx < len(correcciones_titulos) else ""
        titulo = (titulo or "").strip()
        file = correcciones_imagenes[idx] if idx < len(correcciones_imagenes) else None
        img_rel_path = None
        
        if file and file.filename and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            ts = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")
            name, ext = os.path.splitext(filename)
            filename = f"{name}_{ts}{ext}"
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(file_path)
            img_rel_path = f"/uploads/{filename}"
        
        if desc or img_rel_path or titulo:
            saved_correcciones.append({
                "titulo": titulo if titulo else f"Corrección {idx + 1}",
                "descripcion": desc, 
                "imagen": img_rel_path
            })

    ts_base = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"reporte_mantenimiento_{ts_base}"
    docx_path = os.path.join(app.config["GENERATED_FOLDER"], base_name + ".docx")

    # --- Generar DOCX con estilo corporativo CORREGIDO ---
    try:
        document = Document()

        # CREAR ENCABEZADO (método alternativo más confiable)
        # Crear tabla de encabezado en el cuerpo del documento
        header_table = document.add_table(rows=1, cols=3)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar para repetir como encabezado en todas las páginas
        header_table.rows[0]._tr.get_or_add_trPr().append(
            OxmlElement('w:tblHeader')
        )
        
        # Configurar anchos de columnas - IGUALAR CON TABLA DE DATOS
        header_table.columns[0].width = Inches(1.5)  # Logo - IGUAL QUE DATOS
        header_table.columns[1].width = Inches(2.0)  # Título - IGUAL QUE DATOS  
        header_table.columns[2].width = Inches(1.0)  # Versión/Fecha - AJUSTADO
        
        cells = header_table.rows[0].cells
        
        # Configurar altura de fila más compacta
        header_table.rows[0].height = Inches(0.8)
        
        # Configurar bordes y alineación vertical para todas las celdas
        for cell in cells:
            add_cell_borders(cell)
            set_cell_vertical_alignment(cell, 'center')

        # Celda izquierda: Logo NAVITRANS
        left_para = cells[0].paragraphs[0]
        left_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Intentar cargar logo, si no existe usar texto
        logo_path = os.path.join(app.root_path, "static", "img", "logo.png")
        if os.path.exists(logo_path):
            try:
                left_para.add_run().add_picture(logo_path, height=Inches(0.6))
            except:
                # Fallback a texto si hay error con la imagen
                fallback_run = left_para.add_run("NAVITRANS\nMantenimiento")
                set_font_cambria_size(fallback_run, 9)
                fallback_run.bold = True
        else:
            # Texto como fallback si no hay logo
            fallback_run = left_para.add_run("NAVITRANS\nMantenimiento")
            set_font_cambria_size(fallback_run, 9)
            fallback_run.bold = True

        # Celda central: Título con fondo rojo
        mid_para = cells[1].paragraphs[0]
        mid_run = mid_para.add_run("REPORTE TÉCNICO\nSERVICIO TALLER")
        mid_run.bold = True
        set_font_cambria_size(mid_run, 11)
        mid_run.font.color.rgb = RGBColor(255, 255, 255)  # Texto blanco
        mid_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Aplicar fondo rojo a la celda central
        tc_mid = cells[1]._tc
        tcPr_mid = tc_mid.get_or_add_tcPr()
        shd_mid = OxmlElement('w:shd')
        shd_mid.set(qn('w:val'), 'clear')
        shd_mid.set(qn('w:color'), 'auto')
        shd_mid.set(qn('w:fill'), 'E30613')  # Rojo
        tcPr_mid.append(shd_mid)

        # Celda derecha: Versión y fecha centradas
        right_cell = cells[2]
        # Eliminar párrafo original
        right_cell._element.clear_content()
        
        # Crear tabla anidada para VERSION y FECHA
        right_table = right_cell.add_table(rows=2, cols=1)
        right_table.autofit = True
        
        # Configurar bordes de la tabla anidada y centrado
        for row in right_table.rows:
            row.height = Inches(0.3)
            for cell in row.cells:
                add_cell_borders(cell)
                set_cell_vertical_alignment(cell, 'center')
        
        # Fila 1: VERSIÓN centrada
        version_cell = right_table.rows[0].cells[0]
        version_para = version_cell.paragraphs[0]
        version_run = version_para.add_run("VERSIÓN: 01")
        set_font_cambria_size(version_run, 8)
        version_run.bold = True
        version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Fila 2: FECHA centrada
        fecha_cell = right_table.rows[1].cells[0]
        fecha_para = fecha_cell.paragraphs[0]
        fecha_run = fecha_para.add_run(f"FECHA: {fecha or 'XXXXXXXX'}")
        set_font_cambria_size(fecha_run, 8)
        fecha_run.bold = True
        fecha_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Espacio después del header en el cuerpo del documento
        document.add_paragraph()

        # Helper: create a red strip (table cell background) for section title
        def add_section_title(text):
            t = document.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER  # CENTRAR TABLA
            c = t.rows[0].cells[0]
            # Add run
            run = c.paragraphs[0].add_run(text)
            run.bold = True
            set_font_cambria_size(run, 12)  # 12pt para títulos de sección
            # Apply red fill to the cell (w:shd)
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'E30613')  # red without '#'
            tcPr.append(shd)
            # Set font color to white
            run.font.color.rgb = RGBColor(255, 255, 255)
            return t

        add_section_title("1. DATOS GENERALES")
        
        # Agregar espacio entre el título y la tabla
        document.add_paragraph()
        
        # Crear tabla para datos generales con filas alternadas en gris - TABLA CENTRADA
        datos_table = document.add_table(rows=4, cols=2)
        datos_table.autofit = False  
        
        # CENTRAR TABLA EN LA PÁGINA
        datos_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # REDUCIR significativamente los anchos de columnas
        datos_table.columns[0].width = Inches(1.5)  # Columna etiquetas - REDUCIDO de 2.0 a 1.5
        datos_table.columns[1].width = Inches(2.0)   # Columna valores - REDUCIDO de 2.5 a 2.0
        
        # Datos para la tabla
        datos = [
            ("CLIENTE:", cliente or 'XXXXXXXXXX'),
            ("EQUIPO:", equipo or 'XXXXXXXXXX'),
            ("KILOMETRAJE:", kilometraje or 'XXXXXXXXXX'),
            ("HORAS:", horas or 'XXXXXXXXXX')
        ]
        
        # Colores alternados (gris claro y gris más oscuro)
        colores = ['E6E6E6', 'D0D0D0']  # Gris claro, Gris más oscuro
        
        for i, (etiqueta, valor) in enumerate(datos):
            row = datos_table.rows[i]
            color_fila = colores[i % 2]  # Alterna entre los dos colores
            
            # Celda de etiqueta (columna izquierda)
            celda_etiqueta = row.cells[0]
            set_cell_vertical_alignment(celda_etiqueta, 'center')
            
            # Aplicar color de fondo
            tc_etiqueta = celda_etiqueta._tc
            tcPr_etiqueta = tc_etiqueta.get_or_add_tcPr()
            shd_etiqueta = OxmlElement('w:shd')
            shd_etiqueta.set(qn('w:val'), 'clear')
            shd_etiqueta.set(qn('w:color'), 'auto')
            shd_etiqueta.set(qn('w:fill'), color_fila)
            tcPr_etiqueta.append(shd_etiqueta)
            
            # Texto de etiqueta - CENTRADO TAMBIÉN
            para_etiqueta = celda_etiqueta.paragraphs[0]
            run_etiqueta = para_etiqueta.add_run(etiqueta)
            run_etiqueta.bold = True
            set_font_cambria_size(run_etiqueta, 11)
            para_etiqueta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # CENTRADO
            
            # Celda de valor (columna derecha)
            celda_valor = row.cells[1]
            set_cell_vertical_alignment(celda_valor, 'center')
            
            # Aplicar mismo color de fondo
            tc_valor = celda_valor._tc
            tcPr_valor = tc_valor.get_or_add_tcPr()
            shd_valor = OxmlElement('w:shd')
            shd_valor.set(qn('w:val'), 'clear')
            shd_valor.set(qn('w:color'), 'auto')
            shd_valor.set(qn('w:fill'), color_fila)
            tcPr_valor.append(shd_valor)
            
            # Texto de valor - CENTRADO tanto horizontal como vertical
            para_valor = celda_valor.paragraphs[0]
            run_valor = para_valor.add_run(valor)
            set_font_cambria_size(run_valor, 11)
            para_valor.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Agregar espacio moderado entre secciones
        document.add_paragraph()
        
        add_section_title("2. CONDICIONES")
        for c in condiciones:
            marcado = "☑" if c.get("checked") else "☐"
            cond_para = document.add_paragraph()
            cond_run = cond_para.add_run(f"{marcado} {c.get('text','-')}")
            set_font_cambria_10(cond_run)

        add_section_title("3. CORRECCIONES")
        
        # Agregar espacio entre el título y las correcciones
        document.add_paragraph()
        
        # NUEVA IMPLEMENTACIÓN DE CORRECCIONES CON TABLAS
        for i, corr in enumerate(saved_correcciones, start=1):
            # Espacio antes de cada tabla de corrección (excepto la primera)
            if i > 1:
                document.add_paragraph()
            
            # Crear tabla de 2 filas x 2 columnas para cada corrección
            corr_table = document.add_table(rows=2, cols=2)
            corr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Tamaño intermedio para las tablas de corrección
            corr_table.columns[0].width = Inches(1.75)  # Intermedio para imagen
            corr_table.columns[1].width = Inches(2.25)  # Intermedio para descripción
            
            # FILA 1: Título (combinar columnas)
            title_cell = merge_cells_horizontal(corr_table, 0, 0, 1)
            
            # Configurar título con fondo rojo
            title_para = title_cell.paragraphs[0]
            title_run = title_para.add_run(corr.get("titulo", f"Corrección {i}"))
            title_run.bold = True
            set_font_cambria_size(title_run, 11)
            title_run.font.color.rgb = RGBColor(255, 255, 255)  # Texto blanco
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_cell_vertical_alignment(title_cell, 'center')
            
            # Aplicar fondo rojo al título
            tc_title = title_cell._tc
            tcPr_title = tc_title.get_or_add_tcPr()
            shd_title = OxmlElement('w:shd')
            shd_title.set(qn('w:val'), 'clear')
            shd_title.set(qn('w:color'), 'auto')
            shd_title.set(qn('w:fill'), 'E30613')  # Rojo
            tcPr_title.append(shd_title)
            
            # FILA 2: Imagen en columna 1
            img_cell = corr_table.rows[1].cells[0]
            set_cell_vertical_alignment(img_cell, 'center')
            
            if corr.get("imagen"):
                img_abs = os.path.join(app.root_path, corr["imagen"].lstrip("/"))
                try:
                    img_para = img_cell.paragraphs[0]
                    img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Tamaño intermedio de imagen
                    img_para.add_run().add_picture(img_abs, width=Inches(1.55))
                except Exception:
                    error_para = img_cell.paragraphs[0]
                    error_run = error_para.add_run("(No se pudo insertar la imagen)")
                    set_font_cambria_10(error_run)
                    error_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                # Placeholder si no hay imagen
                no_img_para = img_cell.paragraphs[0]
                no_img_run = no_img_para.add_run("(Sin imagen)")
                set_font_cambria_10(no_img_run)
                no_img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # FILA 2: Descripción en columna 2
            desc_cell = corr_table.rows[1].cells[1]
            set_cell_vertical_alignment(desc_cell, 'center')
            
            # Aplicar fondo gris claro a la celda de descripción
            tc_desc = desc_cell._tc
            tcPr_desc = tc_desc.get_or_add_tcPr()
            shd_desc = OxmlElement('w:shd')
            shd_desc.set(qn('w:val'), 'clear')
            shd_desc.set(qn('w:color'), 'auto')
            shd_desc.set(qn('w:fill'), 'F5F5F5')  # Gris claro
            tcPr_desc.append(shd_desc)
            
            if corr.get("descripcion"):
                desc_para = desc_cell.paragraphs[0]
                desc_run = desc_para.add_run(corr.get("descripcion"))
                set_font_cambria_10(desc_run)
                desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justificar texto
            else:
                # Placeholder si no hay descripción
                no_desc_para = desc_cell.paragraphs[0]
                no_desc_run = no_desc_para.add_run("(Sin descripción)")
                set_font_cambria_10(no_desc_run)
                no_desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Aplicar solo bordes exteriores gruesos
            add_thick_outer_borders_only(corr_table)
            
            # Espacio después de cada tabla de corrección
            document.add_paragraph()

        # Footer with page number field (centered)
        section = document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Crear footer con campo PAGE dinámico
        try:
            # Texto "Página "
            run1 = footer_para.add_run("Página ")
            set_font_cambria_size(run1, 9)
            
            # Campo PAGE dinámico
            run2 = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run2._r.append(fldChar1)
            run2._r.append(instrText)
            run2._r.append(fldChar2)
            set_font_cambria_size(run2, 9)
            
        except Exception:
            # Fallback simple
            footer_run = footer_para.add_run("Página")
            set_font_cambria_size(footer_run, 9)

        document.save(docx_path)
    except Exception as e:
        flash(f"Error generando DOCX: {e}", "danger")
        return redirect(url_for("index"))

    # Result page - solo DOCX por ahora
    docx_filename = os.path.basename(docx_path)
    
    return render_template("result.html", 
                         docx_file=docx_filename, 
                         pdf_file=None)

@app.route("/generated/<path:filename>")
def generated_files(filename):
    return send_from_directory(app.config["GENERATED_FOLDER"], filename, as_attachment=True)

@app.route("/uploads/<path:filename>")
def uploaded_files(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename, as_attachment=False)

@app.route("/health")
def health_check():
    """Health check endpoint para servicios de hosting"""
    return {"status": "ok", "message": "App is running"}

if __name__ == "__main__":
    ensure_dirs()
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_ENV") == "development"
    app.run(host="0.0.0.0", port=port, debug=debug)